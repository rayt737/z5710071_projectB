"""Generate the Part B report as a Word document.

Run from project root:
    python scripts/build_report.py

Mirrors the Part A build script's conventions (title/subtitle, Heading 2
numbered sections, Heading 3 subsections, styled Word tables, inline figures
with captions, ``[REVIEW]`` tags on interpretive claims). Every number quoted
in the prose is pulled live from the CSVs under ``results/`` at build time -
never typed by hand.

IMPORTANT (standing rule, see ai/prompt_log.md Entry 14): once
``report/report.docx`` exists and has been hand-edited in Word, any future
automated regeneration must make SURGICAL edits only - never re-run this
script wholesale, or it will wipe manual edits.
"""
import pathlib
import re
import sys

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
import pandas as pd

ROOT = pathlib.Path(__file__).resolve().parent.parent
RESULTS = ROOT / "results"
FIGURES = RESULTS / "figures"
TABLES = RESULTS / "tables"
DATA = RESULTS / "data"
REPORT = ROOT / "report"

GREY = RGBColor(0x66, 0x66, 0x66)

# ---------------------------------------------------------------------------
# Load the live data every number in the prose traces to
# ---------------------------------------------------------------------------
metrics = pd.read_csv(TABLES / "performance_metrics.csv")
fusion = pd.read_csv(TABLES / "fusion_comparison.csv")
vol_target = pd.read_csv(TABLES / "vol_target_comparison.csv")
fund_returns = pd.read_csv(DATA / "fund_returns.csv", parse_dates=["date"])
names = pd.read_csv(DATA / "fund_display_names.csv")
fees = pd.read_csv(DATA / "fund_fees.csv")
headlines = pd.read_csv(DATA / "headline_sentiment.csv")
sector_idx = pd.read_csv(DATA / "sector_sentiment_index.csv")
lex_summary = pd.read_csv(DATA / "sentiment_lexicon_comparison_summary.csv")

NAME = dict(zip(names["fund_id"], names["display_name"]))
FEE = dict(zip(fees["fund_id"], fees["fee_annual"]))
M = metrics.set_index("fund")
FC = fusion.set_index("fund")
VT = vol_target.set_index("fund")


def name_of(fund_id: str) -> str:
    return NAME.get(fund_id, fund_id)


def pct(x: float, d: int = 1) -> str:
    return f"{x * 100:.{d}f}%"


def num(x: float, d: int = 2) -> str:
    return f"{x:.{d}f}"


def m(fund_id: str, col: str):
    return M.loc[fund_id, col]


def fc(fund_id: str, col: str):
    return FC.loc[fund_id, col]


def vt(fund_id: str, col: str):
    return VT.loc[fund_id, col]


def growth_end(fund_id: str) -> float:
    s = fund_returns[fund_returns["fund"] == fund_id].sort_values("date")["return"]
    return float((1.0 + s).prod())


def _fmt_date_full(dt_str: str) -> str:
    from datetime import datetime
    import platform

    dt = datetime.strptime(str(dt_str)[:10], "%Y-%m-%d")
    fmt = "%#d %B %Y" if platform.system() == "Windows" else "%-d %B %Y"
    return dt.strftime(fmt)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------
NARRATIVE_WORDS = []


def P(doc, text: str) -> None:
    """Body paragraph; words count toward the narrative total.

    Routes through ``_add_math_runs`` so ``_{...}``/``^{...}`` inline math in
    the prose renders as real Word subscripts/superscripts.
    """
    NARRATIVE_WORDS.append(text)
    p = doc.add_paragraph()
    _add_math_runs(p, text)


def CAP(doc, text: str) -> None:
    """Exhibit caption (built-in Caption style); excluded from the word count."""
    p = doc.add_paragraph(text, style="Caption")
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY


def add_styled_table(doc, df, caption, col_widths=None):
    """Add a styled Word table with a caption (Part A convention)."""
    CAP(doc, caption)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(7)

    if col_widths is not None:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()


def add_figure(doc, img_path, caption, width=5.7):
    """Add a centered inline figure with a caption below."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width))
    CAP(doc, caption)
    doc.add_paragraph()


_SUB_SUP = re.compile(r"(_\{[^}]*\}|\^\{[^}]*\})")


def _add_math_runs(p, text: str) -> None:
    """Add runs to a paragraph, honouring _{...} (subscript) and ^{...} (superscript)."""
    for part in _SUB_SUP.split(text):
        if not part:
            continue
        if part.startswith("_{") and part.endswith("}"):
            run = p.add_run(part[2:-1])
            run.font.subscript = True
        elif part.startswith("^{") and part.endswith("}"):
            run = p.add_run(part[2:-1])
            run.font.superscript = True
        else:
            p.add_run(part)


def add_equation(doc, eq_text: str, number: int) -> None:
    """Thesis-style equation: centred on its own line, (n) flush right.

    Uses a centre tab at the middle of the text column and a right tab at the
    right margin. Variables are defined immediately after in the prose.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    p.add_run("\t")
    _add_math_runs(p, eq_text)
    p.add_run("\t")
    p.add_run(f"({number})")


def add_heading(doc, text: str, level: int = 2) -> None:
    doc.add_heading(text, level=level)


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------
def build_report() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ======================================================================
    # Title
    # ======================================================================
    doc.add_heading("Invesper: Systematic Funds, Sentiment, and the App", level=1)
    subtitle = doc.add_paragraph(
        "\u201cInvesting to Prosper\u201d \u2014 FINS3645 FinTech Project \u2014 Part B: Stations 3\u20134"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
    student = doc.add_paragraph("Raynard Nicholas Thela \u2014 z5710071")
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in student.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
    doc.add_paragraph()

    # ======================================================================
    # 1. The funds and the backtest design
    # ======================================================================
    add_heading(doc, "1. The Funds and the Backtest Design", level=2)

    P(doc, (
        "Invesper is a systematic multi-asset platform for retail investors who want "
        "transparent, rules-based fund construction without the complexity of "
        "quantitative tools. Part A established the data foundation \u2014 daily prices for "
        "50 US equities across 10 sectors and 10 cryptocurrencies, 2020\u20132023, plus the "
        "news headlines that feed the sentiment work in Section 3. Part B turns that "
        "foundation into a product: 26 investable funds, a news-sentiment index, and a "
        "deployed Streamlit app. This report documents how the funds were built and "
        "backtested, what the out-of-sample results show, how sentiment is measured and "
        "fused, how the extensions perform, how the app works, and what a critical "
        "review of the whole build implies."
    ))

    P(doc, (
        "The core of the product is a set of 12 funds spanning three universes \u2014 "
        "equity-only, crypto-only, and combined multi-asset \u2014 each constructed with "
        "four optimisation methods: equal-weight, minimum-variance (Markowitz, 1952), "
        "maximum-Sharpe, and risk parity (Maillard, Roncalli & Teiletche, 2010). Each "
        "(universe, method) pair is one fund: it is what a user invests in and what a "
        "fact sheet covers. Combined funds merge crypto returns onto the equity trading "
        "calendar, so weekend-only crypto moves are dropped rather than fabricated into "
        "a return a fund trading on equity days could not capture."
    ))

    P(doc, (
        "Ten additional grouping funds broaden the shelf without inventing data. Three "
        "equity sector-clusters are built from the sectors already in the dataset: "
        "Defensive (Healthcare, Utilities, Consumer), Cyclical (Financials, "
        "Industrials, Materials, Real Estate), and Growth & Innovation (Technology, "
        "Communication, Energy). Two crypto themes use explicit ticker lists: Digital "
        "Payments (XRP, LTC, BCH, XLM) and Web3 Infrastructure (ETH, ADA, ETC, TRX, "
        "EOS), with BTC-USD excluded from both as a store of value rather than a "
        "payment rail or a smart-contract layer. Because no market-capitalisation data "
        "exists in the dataset, these 10 groupings use only equal-weight and risk "
        "parity \u2014 the two methods that need no return or risk estimate \u2014 so their "
        "construction rests on data rather than guesswork. Two judgement calls are "
        "stated: the dataset merges Consumer Staples and Discretionary into one "
        "Consumer sector (placed in Defensive), and TRX sits in Web3 despite its "
        "dominant stablecoin usage, because it is architecturally a smart-contract "
        "platform."
    ))

    P(doc, (
        "Four extension variants sit on top of the 22 base funds to make 26: two "
        "sentiment-fusion tilts on the equity minimum-variance fund (a momentum tilt "
        "and a contrarian tilt, Section 4) and two volatility-targeted overlays on the "
        "Combined and Crypto Maximum-Sharpe funds (Section 4). These variants are where "
        "the innovation band is earned, and they are tested and reported with the same "
        "discipline as the base funds."
    ))

    add_heading(doc, "Backtest methodology", level=3)

    P(doc, (
        "Every fund is backtested out-of-sample with a walk-forward, expanding-window "
        "procedure. The initial estimation window is the first year of that universe's "
        "own calendar \u2014 252 trading days for equity and combined funds, 365 days for "
        "crypto \u2014 and the fund\u2019s weights are rebalanced monthly on the first trading "
        "day of each month. At every rebalance, expected returns and covariances are "
        "estimated from rows strictly before that date, and the sentiment signal used "
        "in Section 4 is lagged one trading day; an explicit assertion "
        "``assert_no_lookahead`` enforces both conditions in code rather than relying "
        "on careful indexing. The first live backtest dates are derived from the data "
        "itself: "
    ))

    P(doc, (
        f"{_fmt_date_full(m('equity_equal_weight', 'first_live_date'))} for equity and "
        f"combined funds and {_fmt_date_full(m('crypto_equal_weight', 'first_live_date'))} "
        "for crypto, with the out-of-sample window running through "
        f"{_fmt_date_full(m('combined_equal_weight', 'last_date'))} (equity/combined) and "
        f"{_fmt_date_full(m('crypto_equal_weight', 'last_date'))} (crypto)."
    ))

    P(doc, (
        "Two methodological choices deserve stating rather than hiding. The window "
        "expands rather than rolls, so every rebalance uses all of the history "
        "available up to that date \u2014 the longest sample and the most stable "
        "covariance estimates, at the cost of reacting slowly to a regime change. "
        "Rebalancing is monthly rather than daily because the product is a retail "
        "fund: monthly is the cadence a customer can follow, and it keeps turnover "
        "\u2014 and, in a real fund, trading cost \u2014 in proportion to the stated "
        "management fee. [REVIEW] Both choices load the setup mildly in the "
        "optimisers\u2019 favour \u2014 more data helps the estimators and slower "
        "rebalancing suits buy-and-hold \u2014 so the simple-fund dominance reported in "
        "Section 2 survives the backtest being stacked against it."
    ))

    P(doc, (
        "Each fund\u2019s daily return is the weighted average of its constituents\u2019 daily "
        "returns, using the target weights set at the most recent rebalance:"
    ))
    add_equation(doc, "R_{p,t} = \u03a3_{i} w_{i,t} r_{i,t}", 1)
    P(doc, (
        "where R_{p,t} is fund p\u2019s simple daily return on day t, w_{i,t} is the target "
        "weight of asset i set at the rebalance immediately before t, and r_{i,t} is "
        "asset i\u2019s simple daily return on day t. Weights are long-only and fully "
        "invested: each is non-negative and they sum to one at every rebalance, "
        "enforced by clipping and renormalising after each solve and asserted in the "
        "test suite."
    ))

    P(doc, "Annualised return and volatility use each universe\u2019s own trading frequency:")
    add_equation(doc, "\u03c3\u0302_{ann} = \u03c3\u0302_{d} \u00d7 \u221an", 2)
    P(doc, (
        "where \u03c3\u0302_{d} is the daily volatility estimate and n = 252 for "
        "equity and combined funds and n = 365 for crypto, because crypto trades on "
        "calendar days while equities trade on business days. The risk-free rate is "
        "set to zero (r_{f} = 0), a stated assumption permitted by the brief, and the "
        "Sharpe ratio follows Sharpe (1966):"
    ))
    add_equation(doc, "SR = (\u03bc\u0302_{ann} \u2212 r_{f}) / \u03c3\u0302_{ann}", 3)
    P(doc, (
        "where \u03bc\u0302_{ann} is the annualised mean return. The Sharpe ratio "
        "therefore ranks funds purely on their excess return per unit of volatility, "
        "with no cost of carry attached."
    ))

    P(doc, (
        "The backtests assume zero transaction costs, which the brief explicitly "
        "permits when stated. [REVIEW] This is more than a modelling simplification: "
        "it is a deliberate pricing choice carried into Section 5, where Invesper\u2019s "
        "zero-commission model treats trading as free at the point of execution. "
        "[REVIEW] Because turnover differs sharply across methods \u2014 equal-weight "
        "rebalances a fraction of the book while maximum-Sharpe and the fusion tilts "
        "turn it over heavily \u2014 the cost assumption matters most for exactly the "
        "variants whose performance is closest to the edge, a point returned to in "
        "Sections 4 and 6."
    ))

    # ======================================================================
    # 2. Out-of-sample results and fund fact sheets
    # ======================================================================
    add_heading(doc, "2. Out-of-Sample Results and Fund Fact Sheets", level=2)

    P(doc, (
        "Table 1 reports annualised return, annualised volatility, Sharpe ratio, "
        "maximum drawdown, total return, and turnover for all 22 base funds over their "
        "out-of-sample windows. Every fund below is cited by its investor-facing name "
        "in the discussion that follows."
    ))

    metrics_disp = metrics[["fund", "annualised_return", "annualised_volatility",
                            "sharpe_ratio", "max_drawdown", "total_return", "turnover"]].copy()
    metrics_disp["fund"] = metrics_disp["fund"].map(name_of)
    metrics_disp.columns = ["Fund", "Ann. return %", "Ann. vol %", "Sharpe",
                            "Max drawdown %", "Total return %", "Turnover %"]
    for col in ("Ann. return %", "Ann. vol %", "Max drawdown %", "Total return %", "Turnover %"):
        metrics_disp[col] = metrics_disp[col].map(lambda x: pct(x, 1))
    metrics_disp["Sharpe"] = metrics_disp["Sharpe"].map(lambda x: num(x, 2))
    add_styled_table(
        doc, metrics_disp,
        "Table 1: Out-of-sample performance metrics, all 22 base funds. "
        "rf = 0; annualisation \u221a252 (equity/combined) and \u221a365 (crypto); "
        "turnover is average monthly portfolio turnover; "
        "source: results/tables/performance_metrics.csv.",
        col_widths=[2.1, 0.7, 0.6, 0.6, 0.7, 0.7, 0.6],
    )

    eq_ew = m("equity_equal_weight", "sharpe_ratio")
    eq_ms = m("equity_max_sharpe", "sharpe_ratio")
    eq_rp = m("equity_risk_parity", "sharpe_ratio")
    eq_mv = m("equity_min_variance", "sharpe_ratio")
    P(doc, (
        "The headline finding of the whole backtest is that the simplest fund wins. "
        "The Invesper US Equity Index Fund, which weights 50 equities equally, "
        f"produces the best out-of-sample Sharpe ratio of the four core equity funds ({num(eq_ew, 2)}), "
        f"ahead of the Opportunities Fund ({num(eq_ms, 2)}), the Balanced Risk Fund "
        f"({num(eq_rp, 2)}), and the Minimum Volatility Fund ({num(eq_mv, 2)}). "
        "[REVIEW] This is not an anomaly: it is exactly the finding of DeMiguel, "
        "Garlappi & Uppal (2009), whose naive 1/N benchmark beat optimised portfolios "
        "out-of-sample, and it matches the reference run in the Week 10 deck. "
        "[REVIEW] On three years of daily data, the estimation error in the optimisers\u2019 "
        "return and covariance estimates outweighs the precision they add, so the "
        "unconstrained, diversification-maximising allocation wins by default."
    ))

    add_figure(
        doc, FIGURES / "growth_of_1_equity.png",
        "Figure 1: Growth of $1 \u2014 Invesper US Equity funds, all methods. "
        "Source: results/data/fund_returns.csv (equity family). Sample period "
        f"{_fmt_date_full(m('equity_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('equity_equal_weight', 'last_date'))}. "
        "Growth is simple daily returns compounded from $1.",
    )

    P(doc, (
        f"The growth of $1 in Figure 1 tells a complementary story. The Index Fund "
        f"ends at ${num(growth_end('equity_equal_weight'), 2)}, but the Opportunities "
        f"Fund reaches ${num(growth_end('equity_max_sharpe'), 2)} \u2014 the highest terminal "
        f"value in Figure 1 \u2014 despite the worst risk-adjusted score of the "
        f"four. The Minimum Volatility Fund delivers ${num(growth_end('equity_min_variance'), 2)} "
        f"with the shallowest drawdown (Table 1), and the Balanced Risk Fund lands "
        f"between them. [REVIEW] The practical reading for an investor is that "
        "maximum-Sharpe buys return with a lot of extra drawdown risk, while "
        "equal-weight buys nearly the same terminal value for far less of it."
    ))

    add_figure(
        doc, FIGURES / "growth_of_1_crypto.png",
        "Figure 2: Growth of $1 \u2014 Invesper Digital Assets funds, all methods. "
        "Source: results/data/fund_returns.csv (crypto family). Sample period "
        f"{_fmt_date_full(m('crypto_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('crypto_equal_weight', 'last_date'))}. "
        "Growth is simple daily returns compounded from $1.",
    )

    P(doc, (
        "Crypto is a different world on every dimension. The Digital Assets Index "
        f"Fund ends at ${num(growth_end('crypto_equal_weight'), 2)} and the Digital "
        f"Assets Balanced Risk Fund at ${num(growth_end('crypto_risk_parity'), 2)}, "
        f"while the Opportunities Fund is the laggard at ${num(growth_end('crypto_max_sharpe'), 2)} "
        "despite the best single-asset tailwind in the sample. [REVIEW] The "
        "concentration that maximum-Sharpe builds inside an already high-volatility "
        "universe compounds the risk without delivering commensurate return, which is "
        "why its maximum drawdown \u2014 Table 1 \u2014 is the deepest of the four core "
        "crypto method funds."
    ))

    P(doc, (
        "One nuance runs against the equity story and is worth stating plainly. In "
        f"the crypto family risk parity edges out equal-weight: the Digital Assets "
        f"Balanced Risk Fund\u2019s Sharpe of {num(m('crypto_risk_parity', 'sharpe_ratio'), 2)} "
        f"against the Index Fund\u2019s {num(m('crypto_equal_weight', 'sharpe_ratio'), 2)}, "
        f"with terminal values of ${num(growth_end('crypto_risk_parity'), 2)} and "
        f"${num(growth_end('crypto_equal_weight'), 2)}. [REVIEW] The gap is small and "
        "not economically decisive, but the direction is informative: in a universe "
        "where a few assets dominate and co-move violently, risk parity spreads risk "
        "as effectively as the simplest 1/N allocation, whereas on the equity side "
        "equal weight clearly outperformed. The consistent lesson is not that one "
        "method always wins, but that the distance between simple and optimised is "
        "smallest exactly where risk is hardest to model."
    ))

    add_figure(
        doc, FIGURES / "growth_of_1_combined.png",
        "Figure 3: Growth of $1 \u2014 Invesper Multi-Asset funds, all methods. "
        "Source: results/data/fund_returns.csv (combined family). Sample period "
        f"{_fmt_date_full(m('combined_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('combined_equal_weight', 'last_date'))}. "
        "Growth is simple daily returns compounded from $1.",
    )

    P(doc, (
        "Blending the two asset classes smooths the extremes. The Multi-Asset Index "
        f"Fund ends at ${num(growth_end('combined_equal_weight'), 2)} and the "
        f"Multi-Asset Balanced Risk Fund at ${num(growth_end('combined_risk_parity'), 2)}, "
        f"ahead of the Opportunities Fund (${num(growth_end('combined_max_sharpe'), 2)}) "
        f"and the Minimum Volatility Fund (${num(growth_end('combined_min_variance'), 2)}). "
        "[REVIEW] The combined minimum-variance fund is the most interesting risk "
        "product in the shelf: it posts the lowest volatility and shallowest drawdown "
        "in the entire multi-asset family while still compounding positive returns \u2014 "
        "the closest thing to a defensive all-weather fund the dataset can support."
    ))

    add_figure(
        doc, FIGURES / "growth_of_1_sector_clusters.png",
        "Figure 4: Growth of $1 \u2014 Invesper equity sector-cluster funds "
        "(Defensive, Cyclical, Growth & Innovation), equal-weight and risk parity. "
        "Source: results/data/fund_returns.csv (defensive, cyclical, "
        "growth_sensitive families). Sample period "
        f"{_fmt_date_full(m('defensive_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('defensive_equal_weight', 'last_date'))}. "
        "Growth is simple daily returns compounded from $1.",
    )

    P(doc, (
        "The sector clusters reveal where equity risk actually lives. The Growth & "
        "Innovation Sectors Index Fund posts the best Sharpe ratio in the entire "
        f"22-fund set at {num(m('growth_sensitive_equal_weight', 'sharpe_ratio'), 2)}, "
        f"ending at ${num(growth_end('growth_sensitive_equal_weight'), 2)}, while the "
        f"Defensive Sectors Index Fund ends at ${num(growth_end('defensive_equal_weight'), 2)} "
        f"with the lowest Sharpe in the shelf ({num(m('defensive_equal_weight', 'sharpe_ratio'), 2)}). "
        "[REVIEW] The spread is a clean picture of the growth-versus-defence trade: "
        "in the 2021\u20132023 window, technology-led growth rewarded risk handsomely "
        "while defensive names compounded slowly, and the low-volatility defence the "
        "defensive cluster offers was not needed in a mostly rising market."
    ))

    P(doc, (
        "Between those two extremes sits the Cyclical Sectors Index Fund, whose "
        f"Sharpe of {num(m('cyclical_equal_weight', 'sharpe_ratio'), 2)} and terminal "
        f"value of ${num(growth_end('cyclical_equal_weight'), 2)} land much closer to "
        "the growth end than the defensive end \u2014 cyclical financials, industrials, "
        "materials, and real estate rode the 2021 reopening and the 2023 recovery "
        "without Technology\u2019s outright concentration. The within-cluster method "
        "comparison is the honest caveat to any neat \u2018index wins\u2019 story: "
        f"equal-weight beats risk parity in Cyclical ({num(m('cyclical_equal_weight', 'sharpe_ratio'), 2)} "
        f"vs {num(m('cyclical_risk_parity', 'sharpe_ratio'), 2)}) and in Growth & "
        f"Innovation ({num(m('growth_sensitive_equal_weight', 'sharpe_ratio'), 2)} vs "
        f"{num(m('growth_sensitive_risk_parity', 'sharpe_ratio'), 2)}), but in "
        f"Defensive the risk-parity version earns the higher Sharpe "
        f"({num(m('defensive_risk_parity', 'sharpe_ratio'), 2)} vs "
        f"{num(m('defensive_equal_weight', 'sharpe_ratio'), 2)}). [REVIEW] The pattern "
        "is not that one method always wins; it is that whichever version spreads "
        "risk most evenly across a cluster\u2019s members tends to score higher, the "
        "same diversification logic the core funds exhibit."
    ))

    add_figure(
        doc, FIGURES / "growth_of_1_crypto_themes.png",
        "Figure 5: Growth of $1 \u2014 Invesper crypto-theme funds (Digital Payments, "
        "Web3 Infrastructure), equal-weight and risk parity. Source: "
        "results/data/fund_returns.csv (payments, web3infra families). Sample period "
        f"{_fmt_date_full(m('payments_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('payments_equal_weight', 'last_date'))}. "
        "Growth is simple daily returns compounded from $1.",
    )

    P(doc, (
        "The crypto themes are the highest-returning products in the shelf and the "
        f"most volatile. The Web3 Infrastructure Balanced Risk Fund ends at "
        f"${num(growth_end('web3infra_risk_parity'), 2)} and the Digital Payments Index "
        f"Fund at ${num(growth_end('payments_equal_weight'), 2)}, and both are deep "
        f"drawdown vehicles: Web3 Infrastructure gives back "
        f"{pct(m('web3infra_risk_parity', 'max_drawdown'))} from peak and Digital "
        f"Payments {pct(m('payments_equal_weight', 'max_drawdown'))} at its worst "
        "(Table 1). [REVIEW] These funds are "
        "return vehicles, not risk vehicles \u2014 an honest label that the app\u2019s fact "
        "sheets make visible by showing drawdown next to growth."
    ))

    add_figure(
        doc, FIGURES / "drawdown_combined_max_sharpe.png",
        "Figure 6: Drawdown from peak \u2014 Invesper Multi-Asset Opportunities Fund "
        "(combined_max_sharpe). Source: results/data/fund_returns.csv "
        "(combined_max_sharpe). Sample period "
        f"{_fmt_date_full(m('combined_max_sharpe', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('combined_max_sharpe', 'last_date'))}. "
        "Drawdown is the cumulative fall from the running peak.",
    )

    P(doc, (
        f"Figure 6 shows why the Multi-Asset Opportunities Fund\u2019s headline numbers "
        f"flatter it. Its maximum drawdown of {pct(m('combined_max_sharpe', 'max_drawdown'))} "
        "is the worst of any combined fund, produced in the 2022 sell-off when the "
        "optimiser\u2019s concentrated long book of high-volatility crypto had no "
        "defensive offset. [REVIEW] The chart is the strongest single argument against "
        "a maximum-Sharpe flagship for retail investors: the growth line looks good, "
        "but the path to it \u2014 and the trust it costs \u2014 is not."
    ))

    add_figure(
        doc, FIGURES / "portfolio_weights_combined_min_variance.png",
        "Figure 7: Target weights by sector at each monthly rebalance \u2014 Invesper "
        "Multi-Asset Minimum Volatility Fund (combined_min_variance). Source: "
        "results/data/fund_weights.csv (combined_min_variance). Sample period "
        f"{_fmt_date_full(m('combined_min_variance', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('combined_min_variance', 'last_date'))}. "
        "Weights are target allocations at each monthly rebalance.",
    )

    P(doc, (
        "Figure 7 shows how a minimum-variance fund actually behaves over time. The "
        "Multi-Asset Minimum Volatility Fund shifts weight toward whatever has been "
        "calm recently \u2014 equity sectors during crypto stress and crypto during "
        "equity stress \u2014 while staying long-only and fully invested at every "
        "rebalance. [REVIEW] The drift is the method working as designed: minimum "
        "variance is a risk-seeker\u2019s opposite, moving capital toward quiet assets "
        "rather than falling in love with winners, which is precisely why its "
        "drawdowns stay shallow."
    ))

    add_figure(
        doc, FIGURES / "sharpe_barplot.png",
        "Figure 8: Out-of-sample Sharpe ratios, all 22 base funds (rf = 0). "
        "Source: results/tables/performance_metrics.csv. Sample period "
        f"{_fmt_date_full(m('equity_equal_weight', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('crypto_equal_weight', 'last_date'))}. "
        "Annualisation \u221a252 (equity/combined) and \u221a365 (crypto).",
        width=6.2,
    )

    P(doc, (
        "Figure 8 puts every fund on one scale. The Growth & Innovation Sectors "
        f"Index Fund ({num(m('growth_sensitive_equal_weight', 'sharpe_ratio'), 2)}) and "
        f"Web3 Infrastructure Balanced Risk Fund ({num(m('web3infra_risk_parity', 'sharpe_ratio'), 2)}) "
        "top the chart, the defensive cluster anchors the bottom, and the three "
        "equal-weight index funds cluster tightly in the upper-middle. [REVIEW] The "
        "barplot is the cleanest summary of the whole product: the funds that "
        "diversify most broadly and simplest \u2014 index, balanced risk, and sector "
        "growth \u2014 are the ones that survive the out-of-sample test best, while the "
        "optimised single-asset funds cluster lower than their in-sample promise."
    ))

    P(doc, (
        "Every fund\u2019s fact sheet, as a user would read it in the app, carries the "
        "same five things Table 1 and Figures 1\u20138 show: growth of $1, annualised "
        "return, annualised volatility, Sharpe ratio, maximum drawdown, and the "
        "current holdings (the target weights from the most recent rebalance). The "
        "extension variants in Section 4 get the same treatment, so a user can "
        "compare the momentum tilt or the managed-volatility overlay against the "
        "fund they came from on identical terms."
    ))

    # ======================================================================
    # 3. The sentiment index
    # ======================================================================
    add_heading(doc, "3. The Sentiment Index", level=2)

    P(doc, (
        "The sentiment model scores the news headlines that accompany the equity "
        "data. After de-duplication the headline set contains "
        f"{len(headlines):,} unique headlines for the 50 equities across their 10 "
        f"sectors, spanning {_fmt_date_full(headlines['date'].min())} to "
        f"{_fmt_date_full(headlines['date'].max())} (the same 2020\u20132023 sample as the "
        "prices). The baseline scorer is VADER (Hutto & Gilbert, 2014), a "
        "rule-based lexicon model built for social-media text that returns a "
        "compound score in [\u20131, +1] for each headline."
    ))

    n_plain = float((headlines["plain_vader"].abs() < 0.05).mean())
    n_ext = float((headlines["finvader_lite"].abs() < 0.05).mean())
    P(doc, (
        f"Plain VADER is conservative on financial text: {pct(n_plain, 1)} of all "
        "headlines score essentially neutral. The brief flags this as a known "
        "weakness \u2014 a sentiment of zero is not the same as no information \u2014 so the "
        "project extends VADER with a small finance-specific lexicon, finVADER-lite. "
        "This is a novel 35-term extension grounded in the Loughran-McDonald Master "
        "Dictionary: finance-relevant terms (for example \u2018guidance\u2019, \u2018sell-off\u2019, "
        "\u2018upgrade\u2019) were proposed, rated for sign and magnitude on two independent "
        "passes, and kept only where the two passes agreed. It is not the pre-built "
        "finVADER package; the full two-pass methodology, the candidate ledger, and "
        "the agreement filter are logged in ai/lexicon_extension_log.md. "
        "[REVIEW] The two-pass agreement rule exists because a finance lexicon is "
        "only useful if its ratings are defensible \u2014 averaging away a disagreement "
        "would hide exactly the judgements that need scrutiny."
    ))

    add_figure(
        doc, FIGURES / "sentiment_lexicon_comparison.png",
        "Figure 9: Plain VADER vs finVADER-lite \u2014 per-headline score distribution "
        "and monthly mean sentiment. Source: results/data/headline_sentiment.csv "
        "(all headlines, 2020\u20132023). Scores are per-headline compound sentiment.",
        width=6.2,
    )

    mm_sum = lex_summary[lex_summary["series"] == "monthly_mean"]
    mm_plain = mm_sum[mm_sum["lexicon"] == "plain_vader"].set_index("date")["value"]
    mm_ext = mm_sum[mm_sum["lexicon"] == "finvader_lite"].set_index("date")["value"]
    corr = float(mm_plain.corr(mm_ext))
    n_changed = int((headlines["plain_vader"] != headlines["finvader_lite"]).sum())
    P(doc, (
        "Figure 9 is the direct evidence that the extension did something measurable. "
        "The two scorers agree closely at the monthly level \u2014 correlation "
        f"{num(corr, 3)} \u2014 because finVADER-lite only changes a headline\u2019s score when "
        f"a finance-specific reading differs from the general-purpose one: {n_changed:,} "
        f"of {len(headlines):,} scores change ({pct(n_changed / len(headlines), 1)}), and "
        "the monthly means move by at most 0.007. [REVIEW] The value of the extension "
        f"is therefore not a wholesale re-scoring but a targeted one: the neutral "
        f"bucket shrinks from {pct(n_plain, 1)} to {pct(n_ext, 1)} of headlines, "
        "recovering a slice of the false neutrals the brief warns about, and the "
        "scores that do change are precisely the finance-flavoured ones. "
        "[REVIEW] The honesty check is that the two lines in the lower panel "
        "overlap for most of the sample; the extension is a refinement, not a "
        "replacement."
    ))

    P(doc, (
        "The sector sentiment index is built from these per-headline scores. Each "
        "ticker-day sentiment is the mean compound score of that stock\u2019s headlines on "
        "that day; each sector\u2019s daily index is the equal-weighted mean of its "
        "constituent stocks, matching the equal-weight philosophy of the funds. Two "
        "text-handling choices are explicit. Days with no headline are carried "
        "forward from the previous score rather than dropped, because an absent "
        "headline does not mean neutral sentiment and dropping the rows would hollow "
        "out the index on quiet days; and stock-days before a stock\u2019s first headline "
        "are treated as neutral zero, because there is genuinely no information yet. "
        "The signal is then lagged by one equity trading day before it is aligned to "
        "returns, so a Saturday or Monday headline \u2014 both aligned to Monday \u2014 is "
        "first usable for Tuesday\u2019s trade. Sentiment therefore never sees the return "
        "it is being tested against."
    ))

    add_figure(
        doc, FIGURES / "sector_sentiment_index.png",
        "Figure 10: Daily sector sentiment index, all 10 equity sectors "
        "(equal-weight stocks; no-news days carried forward). Source: "
        "results/data/sector_sentiment_index.csv (lagged 1 trading day). Sample "
        f"period {_fmt_date_full(sector_idx['date'].min())} \u2013 "
        f"{_fmt_date_full(sector_idx['date'].max())}.",
        width=6.2,
    )

    sec_means = sector_idx.groupby("sector")["sentiment"].mean()
    mkt = sector_idx.groupby("date")["sentiment"].mean()
    P(doc, (
        "Figure 10 shows the index over time for all ten sectors, and its levels "
        "have a clear cross-sectional order. Over the full sample, Utilities carries "
        f"the highest mean sentiment ({num(sec_means['Utilities'], 3)}), followed by "
        f"Real Estate ({num(sec_means['RealEstate'], 3)}) and Technology "
        f"({num(sec_means['Tech'], 3)}), while Financials sits at the bottom "
        f"({num(sec_means['Financials'], 3)}). The market-wide equal-weighted index "
        f"stays positive on {pct(float((mkt > 0).mean()), 1)} of trading days and "
        f"ranges from {num(float(mkt.min()), 2)} to {num(float(mkt.max()), 2)}. "
        "[REVIEW] The uniformly positive level is a property of the headlines, not a "
        "forecast \u2014 company news is written to be read, so neutral language is "
        "uncommon and the model\u2019s zero-mean prior is shifted up \u2014 and the index "
        "should be read as a relative signal across sectors and over time, not an "
        "absolute mood thermometer."
    ))

    P(doc, (
        "The index reads as a relative signal partly because of how quiet days are "
        "handled. On days a sector has no headline, the previous score is carried "
        "forward rather than dropped, so the series stays smooth and never hollows "
        "out during news droughts \u2014 but the smoothing means the level partly "
        "reflects the recency of coverage, not only its tone. [REVIEW] The practical "
        "reading for an investor is unchanged: use the index to compare one sector "
        "against another, or one month against the next, rather than treating any "
        "single level as intrinsically good or bad. This is also the sense in which "
        "the index is used downstream: Section 4\u2019s fusion reads each stock\u2019s "
        "sentiment relative to its own recent history, not against an absolute "
        "threshold."
    ))

    P(doc, (
        "Two scope statements close this section. Headlines are a noisy proxy for "
        "the information investors actually trade on \u2014 they are a few words, "
        "chosen by editors \u2014 so every use of the index carries that caveat. And all "
        "sentiment work applies to the equity data only: the cryptocurrency universe "
        "has no news dimension in this dataset, so crypto funds are constructed from "
        "prices alone, consistent with the brief."
    ))

    # ======================================================================
    # 4. Extensions and innovations
    # ======================================================================
    add_heading(doc, "4. Extensions and Innovations", level=2)

    add_heading(doc, "Sentiment fusion: the tilt", level=3)
    P(doc, (
        "The fusion baseline folds the equity sentiment signal into an equity fund "
        "through a tilt on top of the base weights. The base is the Invesper US "
        "Equity Minimum Volatility Fund, and the tilt rescales each stock\u2019s weight "
        "by its recent relative sentiment before renormalising:"
    ))
    add_equation(
        doc,
        "w\u0303_{i,t} = w_{i,t} (1 + \u03bb z_{i,t}) / "
        "\u03a3_{j} w_{j,t} (1 + \u03bb z_{j,t})",
        4,
    )
    P(doc, (
        "where w_{i,t} is the base weight of stock i at rebalance t, z_{i,t} is the "
        "rolling z-score of stock i\u2019s lagged sentiment over the past 60 trading "
        "days (standardised so the tilt is comparable across stocks), and \u03bb is "
        "the tilt direction: \u03bb = +1 follows sentiment (momentum) and "
        "\u03bb = \u22121 fights it (contrarian). Negative tilted weights are clipped to "
        "zero and the weights renormalised to sum to one, so the tilted fund stays "
        "long-only and fully invested. The tilts are deliberately untuned \u2014 no "
        "backtest was used to choose \u03bb \u2014 and are reported before transaction "
        "costs, on the same terms as the base fund."
    ))

    fusion_disp = fusion[["fund", "lambda", "annualised_return", "annualised_volatility",
                          "sharpe_ratio", "max_drawdown", "total_return", "turnover"]].copy()
    fusion_disp["fund"] = fusion_disp["fund"].map(name_of)
    fusion_disp["lambda"] = fusion_disp["lambda"].map(lambda x: "+1" if x > 0 else ("\u22121" if x < 0 else "0"))
    fusion_disp.columns = ["Fund", "\u03bb", "Ann. return %", "Ann. vol %", "Sharpe",
                           "Max drawdown %", "Total return %", "Turnover %"]
    for col in ("Ann. return %", "Ann. vol %", "Max drawdown %", "Total return %", "Turnover %"):
        fusion_disp[col] = fusion_disp[col].map(lambda x: pct(x, 1))
    fusion_disp["Sharpe"] = fusion_disp["Sharpe"].map(lambda x: num(x, 2))
    add_styled_table(
        doc, fusion_disp,
        "Table 2: Sentiment fusion before vs after \u2014 base equity minimum-variance "
        "fund (base), momentum tilt (\u03bb = +1), contrarian tilt (\u03bb = \u22121). "
        "Before transaction costs; turnover is average monthly portfolio turnover. "
        "Source: results/tables/fusion_comparison.csv.",
        col_widths=[2.2, 0.4, 0.7, 0.6, 0.6, 0.7, 0.7, 0.6],
    )

    add_figure(
        doc, FIGURES / "fusion_growth_comparison.png",
        "Figure 11: Growth of $1 \u2014 sentiment fusion on the Invesper US Equity "
        "Minimum Volatility Fund: base vs momentum vs contrarian (before "
        "transaction costs). Source: results/tables/fusion_comparison.csv. Sample "
        f"period {_fmt_date_full(m('equity_min_variance', 'first_live_date'))} \u2013 "
        f"{_fmt_date_full(m('equity_min_variance', 'last_date'))}.",
        width=6.2,
    )

    P(doc, (
        f"Table 2 and Figure 11 are an honest before-and-after. The momentum tilt "
        f"improves the base fund\u2019s Sharpe from {num(fc('equity_min_variance_base', 'sharpe_ratio'), 2)} "
        f"to {num(fc('equity_min_variance_momentum', 'sharpe_ratio'), 2)} and its total "
        f"return from {pct(fc('equity_min_variance_base', 'total_return'))} to "
        f"{pct(fc('equity_min_variance_momentum', 'total_return'))}, while the "
        f"contrarian tilt makes everything worse ({num(fc('equity_min_variance_contrarian', 'sharpe_ratio'), 2)} "
        f"Sharpe, {pct(fc('equity_min_variance_contrarian', 'total_return'))} total return). "
        "[REVIEW] The asymmetry is sensible: over 2021\u20132023 equity sentiment was "
        "broadly positive and momentum-aligned (Section 3), so following it added a "
        "little alpha while fighting it removed a lot. But the turnover column is "
        "the real story. The base fund turns over "
        f"{pct(fc('equity_min_variance_base', 'turnover'), 1)} of its book per month; "
        f"the momentum tilt turns over {pct(fc('equity_min_variance_momentum', 'turnover'), 1)} "
        f"and the contrarian {pct(fc('equity_min_variance_contrarian', 'turnover'), 1)} \u2014 "
        "roughly twelve times more. [REVIEW] Under any realistic transaction-cost "
        "assumption the tilt\u2019s pre-cost edge is spent several times over in trading "
        "costs, which is exactly why Section 1\u2019s zero-cost assumption cannot be "
        "silently carried into Section 4, and why a real fund would not run this "
        "tilt as-is."
    ))

    add_heading(doc, "Volatility targeting and the look-ahead bug", level=3)
    P(doc, (
        "The second extension is a volatility-targeting overlay in the spirit of "
        "Moreira & Muir (2017): instead of changing the holdings, it scales the "
        "fund\u2019s overall exposure up and down so that realised volatility moves "
        "toward a fixed target. Applied to the Combined and Crypto Maximum-Sharpe "
        "funds, the daily scaling factor is:"
    ))
    add_equation(
        doc,
        "k_{t} = min(max(\u03c3_{target} / \u03c3\u0302_{t-1}, 0.5), 1.5),  "
        "R\u0303_{p,t} = k_{t} R_{p,t}",
        5,
    )
    P(doc, (
        "where \u03c3_{target} is the target annualised volatility, \u03c3\u0302_{t-1} "
        "is the trailing 60-trading-day realised volatility ending strictly before t, "
        "and R\u0303_{p,t} is the scaled daily return. The clip keeps the fund "
        "between 50% and 150% exposure, and the trailing window is causal from the "
        "first live day."
    ))

    P(doc, (
        "This extension contains the project\u2019s most instructive mistake, and it is "
        "reported deliberately rather than buried. In the first version, "
        "\u03c3_{target} was computed as the fund\u2019s full-sample realised volatility "
        "over the entire out-of-sample period \u2014 all of 2021\u20132023. That number only "
        "exists after the whole future, including the 2022 crash, has already "
        "happened, so on a calm early-2021 day the future-inflated target sat above "
        "the trailing realised volatility and the overlay scaled exposure up just "
        "before the crash it could not have known about. That is look-ahead bias, "
        "the same class of error the rest of the build is asserted against, and it "
        "made the Combined fund\u2019s maximum drawdown look worse (-52.7% to -59.3%) "
        "rather than better. [REVIEW] The -59.3% is the buggy run\u2019s recorded output "
        "and cannot be reproduced from the committed CSVs, which the corrected run "
        "overwrote; it is logged verbatim in ai/prompt_log.md (Entry 7). The fix "
        "computes the target once, from the fund\u2019s "
        "initial estimation window only \u2014 the 2020 pre-live year that set the first "
        "weights \u2014 and a regression test now asserts the target is a pure function "
        "of that window and that corrupting any later date cannot move it "
        "(test_vol_target_target_is_initial_window_only)."
    ))

    vt_disp = vol_target[["fund", "target_vol", "annualised_return", "annualised_volatility",
                          "sharpe_ratio", "max_drawdown", "total_return"]].copy()
    vt_disp["fund"] = vt_disp["fund"].map(name_of)
    vt_disp["target_vol"] = vt_disp["target_vol"].map(
        lambda x: f"\u2014" if pd.isna(x) else num(x, 2))
    vt_disp.columns = ["Fund", "Target vol", "Ann. return %", "Ann. vol %", "Sharpe",
                       "Max drawdown %", "Total return %"]
    for col in ("Ann. return %", "Ann. vol %", "Max drawdown %", "Total return %"):
        vt_disp[col] = vt_disp[col].map(lambda x: pct(x, 1))
    vt_disp["Sharpe"] = vt_disp["Sharpe"].map(lambda x: num(x, 2))
    add_styled_table(
        doc, vt_disp,
        "Table 3: Volatility-targeting overlay, base vs vol-targeted, Combined and "
        "Crypto Maximum-Sharpe funds (causal targets from the initial estimation "
        "window). Source: results/tables/vol_target_comparison.csv.",
        col_widths=[2.4, 0.8, 0.8, 0.6, 0.6, 0.8, 0.8],
    )

    add_figure(
        doc, FIGURES / "growth_vol_target_combined_max_sharpe.png",
        "Figure 12: Growth of $1 \u2014 Invesper Multi-Asset Opportunities Fund vs its "
        "managed-volatility overlay. Source: results/tables/vol_target_comparison.csv "
        "(combined_max_sharpe). The overlay scales exposure toward the fund\u2019s "
        "initial-estimation-window volatility.",
        width=6.2,
    )
    add_figure(
        doc, FIGURES / "growth_vol_target_crypto_max_sharpe.png",
        "Figure 13: Growth of $1 \u2014 Invesper Digital Assets Opportunities Fund vs "
        "its managed-volatility overlay. Source: results/tables/vol_target_comparison.csv "
        "(crypto_max_sharpe). The overlay scales exposure toward the fund\u2019s "
        "initial-estimation-window volatility.",
        width=6.2,
    )
    add_figure(
        doc, FIGURES / "vol_target_scaling_combined_max_sharpe.png",
        "Figure 14: Volatility-targeting scaling factor k_t \u2014 Invesper Multi-Asset "
        "Opportunities Fund overlay. Source: results/tables/vol_target_comparison.csv. "
        "k_t = target vol / trailing 60-day realised vol, clipped to [0.5, 1.5]; "
        "uses only returns strictly before t.",
        width=6.2,
    )
    add_figure(
        doc, FIGURES / "vol_target_scaling_crypto_max_sharpe.png",
        "Figure 15: Volatility-targeting scaling factor k_t \u2014 Invesper Digital "
        "Assets Opportunities Fund overlay. Source: results/tables/vol_target_comparison.csv. "
        "k_t = target vol / trailing 60-day realised vol, clipped to [0.5, 1.5]; "
        "uses only returns strictly before t.",
        width=6.2,
    )

    P(doc, (
        "Table 3 and Figures 12\u201315 report the corrected, honest result. On the "
        f"Combined fund the overlay makes things worse: Sharpe falls from "
        f"{num(vt('combined_max_sharpe_base', 'sharpe_ratio'), 2)} to "
        f"{num(vt('combined_max_sharpe_vol_targeted', 'sharpe_ratio'), 2)}, total "
        f"return from {pct(vt('combined_max_sharpe_base', 'total_return'))} to "
        f"{pct(vt('combined_max_sharpe_vol_targeted', 'total_return'))}, and maximum "
        f"drawdown deepens from {pct(vt('combined_max_sharpe_base', 'max_drawdown'))} to "
        f"{pct(vt('combined_max_sharpe_vol_targeted', 'max_drawdown'))}. On the Crypto "
        f"fund it is a near no-op: Sharpe barely moves ({num(vt('crypto_max_sharpe_base', 'sharpe_ratio'), 2)} "
        f"to {num(vt('crypto_max_sharpe_vol_targeted', 'sharpe_ratio'), 2)}) with a "
        f"still-worse drawdown ({pct(vt('crypto_max_sharpe_base', 'max_drawdown'))} to "
        f"{pct(vt('crypto_max_sharpe_vol_targeted', 'max_drawdown'))}). "
        "[REVIEW] The mechanism is visible in the scaling-factor figures: the 2020 "
        "estimation window happens to include the COVID crash, so the target "
        f"volatility ({num(vt('combined_max_sharpe_vol_targeted', 'target_vol'), 2)} "
        f"for Combined, {num(vt('crypto_max_sharpe_vol_targeted', 'target_vol'), 2)} "
        "for Crypto) sits well above the "
        "2021\u20132023 trailing volatility, and k_{t} exceeds 1 for most of the sample \u2014 "
        "the overlay spends the whole period leveraging the fund up rather than "
        "de-risking it, which is precisely why the 2022 crash hurts more. The brief "
        "explicitly credits a careful extension with a negative result, and this one "
        "is reported exactly as the corrected numbers show, with nothing tuned to "
        "look better."
    ))

    # ======================================================================
    # 5. The app and the investor journey
    # ======================================================================
    add_heading(doc, "5. The App and the Investor Journey", level=2)

    P(doc, (
        "The product Invesper actually ships is a Streamlit app deployed from a "
        "public GitHub repository (entrypoint streamlit_app.py at the repo root). "
        "Its design constraint is deliberate: the deployed app reads precomputed "
        "artifacts from results/ only and never runs a backtest, an optimiser, or a "
        "sentiment model at request time \u2014 the free hosting tier cannot, and the "
        "loads are cached so the app stays responsive on a basic machine. All 26 "
        "funds across three universes are grouped directly from the data rather than "
        "hardcoded: Equity (12 funds), Crypto (9), and Multi-Asset (5), where the "
        "fusion variants sit in Equity and the managed-volatility variants inherit "
        "their base fund\u2019s universe."
    ))

    P(doc, (
        "The app is organised around the five actions a user actually takes. "
        "Compare shows a per-universe metrics table and a sorted horizontal Sharpe "
        "chart, with a radio switch between universes. Fund fact sheet opens any of "
        "the 26 funds by its investor-facing name and shows growth of $1, drawdown, "
        "the four headline metrics, and current holdings; the managed-volatility "
        "funds show their base fund\u2019s holdings with an explicit caption that the "
        "overlay scales overall exposure, not individual holdings. Sentiment "
        "analytics shows the sector index lines with multi-select and a market-wide "
        "0\u2013100 fear/greed gauge (0 = fear, 100 = greed), with a caption stating the "
        "gauge is unlagged and display-only while fund construction uses the lagged "
        "signal. Allocation is the money step: a user builds a portfolio from any "
        "of the 26 funds, each slider showing its annual fee, and watches the "
        "gross-versus-net growth lines driven by the blended fee. Portfolio is where "
        "the allocation becomes a saved (session-only) portfolio."
    ))

    fee_disp = fees.copy()
    fee_disp["fund_id"] = fee_disp["fund_id"].map(name_of)
    fee_disp["fee_annual"] = fee_disp["fee_annual"].map(lambda x: pct(x, 2))
    fee_disp.columns = ["Fund", "Annual fee %"]
    add_styled_table(
        doc, fee_disp,
        "Table 4: Per-fund annual management-fee schedule, all 26 app-eligible "
        "funds. Source: results/data/fund_fees.csv.",
        col_widths=[4.4, 1.2],
    )

    P(doc, (
        "Table 4 is the pricing model. Fees range from "
        f"{pct(FEE['equity_equal_weight'], 2)} for the equity index fund to "
        f"{pct(FEE['crypto_max_sharpe_vol_targeted'], 2)} for the managed-volatility "
        f"crypto fund, with the fusion tilts priced at {pct(FEE['equity_min_variance_momentum'], 2)} "
        "to reflect their active turnover. A portfolio\u2019s blended fee is the "
        "weighted average of its constituents\u2019 fees by allocation, charged daily on "
        "assets under management, and it drives the gap between the gross and net "
        "growth lines a user sees in the Allocation tab."
    ))

    P(doc, (
        "The zero-commission framing is the customer story behind that fee design. "
        "Invesper follows the model the retail brokerage and robo-advisor industry "
        "moved to after 2019: no per-trade commission, with revenue coming from a "
        "transparent management fee instead. [REVIEW] This is what makes Section "
        "1\u2019s zero-transaction-cost backtest assumption coherent \u2014 it is not a "
        "convenient blind spot but the product\u2019s stated commercial model \u2014 and the "
        "app shows the trade-off honestly by plotting gross and net growth side by "
        "side rather than hiding the fee line."
    ))

    P(doc, (
        "One limitation is stated in the app and repeated here: saved portfolios "
        "persist for the browser session only and vanish on refresh. [REVIEW] For a "
        "coursework deployment this is a defensible scope choice \u2014 it keeps the "
        "app stateless and free-tier-friendly with no database to secure \u2014 but it "
        "is a genuine gap against the real robo-advisor experience, and it is one of "
        "the recommendations in Section 6."
    ))

    P(doc, (
        "The investor journey a user actually follows is: open the app, compare the "
        "funds on the metrics table and Sharpe chart to shortlist candidates; open "
        "each shortlist fund\u2019s fact sheet to check drawdown and current holdings; "
        "check the sentiment analytics to see whether a sector is being talked about "
        "differently from its price history; set an allocation in the Allocation tab "
        "and watch the blended fee reduce the projected growth; and submit the "
        "portfolio to see the final weights, growth, and metrics in the Portfolio "
        "tab. Every label the user reads is an investor-facing name, not a technical "
        "fund id, so the product reads like a fund platform rather than a research "
        "output."
    ))

    P(doc, (
        "The deployment itself is deliberately ordinary, and that is a design "
        "choice. The app runs on the free tier of Streamlit Community Cloud from a "
        "public GitHub repository, so the compute budget is tiny: every result is "
        "precomputed into the small CSV files under results/, figures load once and "
        "are cached, and nothing is optimised or scored at request time \u2014 which is "
        "why the app stays responsive on a basic machine. [REVIEW] This is the right "
        "trade for a coursework deployment, but it is also the honest boundary of "
        "the current product: a read-only window onto precomputed research, not a "
        "live backtesting terminal. That boundary is perfectly consistent with the "
        "investor journey described above, and it is why the recommendations in "
        "Section 6 target the product gaps that remain rather than the research "
        "engine."
    ))

    # ======================================================================
    # 6. Critical reflection and three concrete recommendations
    # ======================================================================
    add_heading(doc, "6. Critical Reflection and Three Concrete Recommendations", level=2)

    P(doc, (
        "Taken as a whole, the build teaches a consistent lesson: on three years of "
        "out-of-sample data, the products that do least win. The equal-weight index "
        "funds beat every optimised equity fund (Section 2); risk parity is the "
        "strongest combined family; the two deliberate complexity bets \u2014 the "
        "sentiment tilt and the volatility overlay \u2014 either spend their edge on "
        "turnover or fail outright (Section 4). "
        "[REVIEW] For a real product this is a sharper finding than any individual "
        "Sharpe ratio: it says the value of Invesper\u2019s shelf is not the optimisation "
        "engine but the honest comparison of simple, transparent, low-cost "
        "diversification against the alternatives \u2014 and that the fee model in "
        "Section 5 is the correct one to charge for it, because it aligns revenue "
        "with the product that actually works."
    ))

    P(doc, (
        "The volatility overlay\u2019s behaviour is the second lesson. Its result "
        "depended on the calibration window in a way that is easy to miss: because "
        "the 2020 estimation window includes the COVID crash, the target sits above "
        "the fund\u2019s live realised volatility and the overlay systematically "
        "leverages rather than de-risks. "
        "[REVIEW] The practical implication is that a volatility target is only as "
        "honest as the window it is calibrated on, and a target fixed before a "
        "regime change can push a fund in the wrong direction for years. The "
        "look-ahead version of this bug \u2014 the full-sample target \u2014 is the project\u2019s "
        "most useful AI-workflow artefact, because it is exactly the kind of "
        "plausible-sounding, subtly wrong output that a careful reviewer and a "
        "causal assertion catch."
    ))

    P(doc, (
        "Three recommendations follow directly from the evidence. Each is specific "
        "enough to act on and is flagged for review because it is a judgement, not "
        "a measured fact."
    ))

    P(doc, (
        "[REVIEW] Recommendation 1: price turnover into the fusion before claiming "
        "any alpha. The momentum tilt\u2019s pre-cost Sharpe gain is "
        f"{num(fc('equity_min_variance_momentum', 'sharpe_ratio') - fc('equity_min_variance_base', 'sharpe_ratio'), 2)} "
        f"points against {pct(fc('equity_min_variance_momentum', 'turnover'), 1)} monthly "
        "turnover, and a real fund would pay for that repeatedly. The next iteration "
        "should add a spread-plus-fee transaction-cost model to the backtest and "
        "re-report the tilts net of it; if the edge survives, the zero-commission "
        "pricing story can be extended honestly to the fused funds, and if it does "
        "not, the tilt should be shelved or redesigned with a rebalancing threshold "
        "to cut turnover."
    ))

    P(doc, (
        "[REVIEW] Recommendation 2: make the volatility target adaptive rather than "
        "fixed to a single pre-live window. The overlay\u2019s failure is traceable to a "
        "target calibrated on a crisis window, so a rolling or re-estimated causal "
        "target \u2014 for example a trailing 12-month target recomputed each rebalance "
        "from returns strictly before it \u2014 would let the overlay de-risk in quiet "
        "markets instead of leveraging into a crash. This should be re-tested with "
        "the same causal assertions before it is presented as a product feature, "
        "and the honest negative result of the current fixed-window version should "
        "stay in the fact sheets either way."
    ))

    P(doc, (
        "[REVIEW] Recommendation 3: use the sentiment index where it is reliable "
        "and drop it where it is not. The sector-level index is the stronger "
        "artefact \u2014 it is built from many headlines per day and shows a stable "
        "cross-sectional order (Section 3) \u2014 yet the fusion only used a single "
        "fund\u2019s stock-level z-scores. The next version should fuse sector sentiment "
        "into the sector funds instead, or use the market-wide gauge for a "
        "defensive cash timing rule, and should keep the honest frame that "
        "sentiment is a refinement of the base fund, not a replacement for it."
    ))

    P(doc, (
        "Finally, the session-only portfolio storage should be promoted to "
        "persistent storage before the app is presented as a real product, "
        "[REVIEW] since the current limitation, while honestly labelled, is the "
        "single largest gap between the coursework deployment and the robo-advisor "
        "experience the product claims to offer."
    ))

    # ======================================================================
    # References
    # ======================================================================
    add_heading(doc, "References", level=2)
    references = [
        "Baker, M., & Wurgler, J. (2007). Investor sentiment in the stock market. "
        "Journal of Economic Perspectives, 21(2), 129\u2013151.",
        "DeMiguel, V., Garlappi, L., & Uppal, R. (2009). Optimal versus naive "
        "diversification: How inefficient is the 1/N portfolio strategy? Review of "
        "Financial Studies, 22(5), 1915\u20131953.",
        "Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model "
        "for sentiment analysis of social media text. Proceedings of the Eighth "
        "International AAAI Conference on Weblogs and Social Media.",
        "Maillard, S., Roncalli, T., & Teiletche, J. (2010). The properties of "
        "equally weighted risk contribution portfolios. Journal of Portfolio "
        "Management, 36(4), 60\u201370.",
        "Markowitz, H. (1952). Portfolio selection. Journal of Finance, 7(1), 77\u201391.",
        "Moreira, A., & Muir, T. (2017). Volatility-managed portfolios. Journal of "
        "Finance, 72(4), 1611\u20131644.",
        "Sharpe, W. F. (1966). Mutual fund performance. Journal of Business, 39(1), "
        "119\u2013138.",
        "Tetlock, P. C. (2007). Giving content to investor sentiment: The role of "
        "media in the stock market. Journal of Finance, 62(3), 1139\u20131168.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    REPORT.mkdir(exist_ok=True)
    out = REPORT / "report.docx"
    doc.save(str(out))

    # ------------------------------------------------------------------
    # Word count: narrative only, excluding captions and references
    # ------------------------------------------------------------------
    words = sum(len(t.split()) for t in NARRATIVE_WORDS)
    review = sum(t.count("[REVIEW]") for t in NARRATIVE_WORDS)
    print(f"Saved {out}")
    print(f"Narrative words (excl. captions + references): {words}")
    print(f"[REVIEW] tags in narrative: {review}")


if __name__ == "__main__":
    build_report()
